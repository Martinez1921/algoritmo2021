from re import S
import openpyxl
from openpyxl import load_workbook
from docx import Document
import mimetypes
import smtplib
import ssl
from email.message import EmailMessage
import os
from os import remove

DIRECCION_DEL_SERVIDOR = "smtp.gmail.com"
PUERTO = 587
DIRECCION_DE_ORIGEN = "proyectalgorit2021@gmail.com"
CONTRASENA = os.getenv('Pasword')

libro = openpyxl.load_workbook('Inventario.xlsx')
hoja_Productos = libro['productos']
hoja_Clientes = libro['clientes']
hoja_Pedidos = libro['pedidos']

Menu_Principal = input("Seleccione la opcion deseada:\n a. Productos\n b. Clientes\n c. Pedidos\n d. Informes\n e. Varios\n f. Salir\n")
while Menu_Principal != "f":
    if Menu_Principal == "a":
        Menu_Productos = input("Seleccione la opcion deseada:\n i.   Agregar producto\n ii.  Editar producto\n iii. Eliminar producto\n iv.  Listar productos\n v.   Enviar cotizacion por correo\n vi.  Salir\n")
        while Menu_Productos != "vi":
            if Menu_Productos == "i":
                # Agregar producto
                hoja_Productos['A1'].value = "Producto"
                hoja_Productos['B1'].value = "Precio"
                hoja_Productos['C1'].value = "Existencia"
                agregar = []
                AgregarDicctionary = {}
                ProductoAgg = input("Ingrese el nombre del producto:\n")
                AgregarDicctionary['producto'] = ProductoAgg
                PrecioAgg = float(input("Ingrese el precio del producto:\n"))
                AgregarDicctionary['precio'] = PrecioAgg
                CantidAgg = int(input("Ingrese la cantidad existente del producto:\n"))
                AgregarDicctionary['existencia'] = CantidAgg
                agregar.append(AgregarDicctionary)
                
                primeraFila = hoja_Productos.max_row + 1
                for produc in agregar:
                    hoja_Productos['A' + str(primeraFila)].value = produc['producto']
                    hoja_Productos['B' + str(primeraFila)].value = produc['precio']
                    hoja_Productos['C' + str(primeraFila)].value = produc['existencia']
                remove("Inventario.xlsx")
                libro.save("Inventario.xlsx")
                Menu_Productos = input("Seleccione la opcion deseada:\n i.   Agregar producto\n ii.  Editar producto\n iii. Eliminar producto\n iv.  Listar productos\n v.   Enviar cotizacion por correo\n vi.  Salir\n")
            if Menu_Productos == "ii":
                # codigo de Editar producto
                productosEdit = []
                diccionarioProductoEdit = {}

                for rows in range(2, hoja_Productos.max_row + 1):
                    productoedit = hoja_Productos["A" + str(rows)].value
                    precioedit = hoja_Productos["B" + str(rows)].value
                    existenciaedit = hoja_Productos["C"+str(rows)].value
                    diccionarioProductoEdit['producto'] = productoedit
                    diccionarioProductoEdit['precio'] = precioedit
                    diccionarioProductoEdit['existencia'] = existenciaedit
                    productosEdit.append(diccionarioProductoEdit)
                    diccionarioProductoEdit = {}
                IteradorEdit = 1
                for productsEdit in productosEdit:
                    print("Linea: "+str(IteradorEdit)+" Producto: ", productsEdit['producto']+" Precio: "+ str(productsEdit['precio'])+ " Existencia: " + str(productsEdit['existencia']))
                    IteradorEdit += 1
                EditarSelect = int(input("Ingrese el numero de linea a editar:  "))
                Edit = EditarSelect + 1
                hoja_Productos['A1'].value = "Producto"
                hoja_Productos['B1'].value = "Precio"
                hoja_Productos['C1'].value = "Existencia"
                agregarEdit = []
                AgregarDicctionaryEdit = {}
                ProductoAggEdit = input("Ingrese el nuevo nombre del producto:  "+ hoja_Productos['A' + str(Edit)].value +"\n")
                AgregarDicctionaryEdit['producto'] = ProductoAggEdit
                PrecioAggEdit = float(input("Ingrese el nuevo precio del producto:  "+ str(hoja_Productos['B' + str(Edit)].value) +"\n"))
                AgregarDicctionaryEdit['precio'] = PrecioAggEdit
                CantidAggEdit = int(input("Ingrese la nuevo cantidad existente del producto:  "+ str(hoja_Productos['C' + str(Edit)].value) +"\n"))
                AgregarDicctionaryEdit['existencia'] = CantidAggEdit
                agregarEdit.append(AgregarDicctionaryEdit)
                
                for producEdit in agregarEdit:
                    hoja_Productos['A' + str(Edit)].value = producEdit['producto']
                    hoja_Productos['B' + str(Edit)].value = producEdit['precio']
                    hoja_Productos['C' + str(Edit)].value = producEdit['existencia']
                remove("Inventario.xlsx")
                libro.save("Inventario.xlsx")
                Menu_Productos = input("Seleccione la opcion deseada:\n i.   Agregar producto\n ii.  Editar producto\n iii. Eliminar producto\n iv.  Listar productos\n v.   Enviar cotizacion por correo\n vi.  Salir\n")
            if Menu_Productos =="iii":
                #Codigo eliminar
                productoseli = []
                diccionarioProductoEli = {}

                for row in range(2, hoja_Productos.max_row + 1):
                    producto = hoja_Productos["A" + str(row)].value
                    precio = hoja_Productos["B" + str(row)].value
                    existencia = hoja_Productos["C"+str(row)].value
                    diccionarioProductoEli['producto'] = producto
                    diccionarioProductoEli['precio'] = precio
                    diccionarioProductoEli['existencia'] = existencia
                    productoseli.append(diccionarioProductoEli)
                    diccionarioProductoEli = {}
                iteraddor = 1
                for productseli in productoseli:
                    print("Linea: "+str(iteraddor)+" Producto: ", productseli['producto']+" Precio: "+ str(productseli['precio'])+ " Existencia: " + str(productseli['existencia']))
                    iteraddor = iteraddor + 1
                LineaDelete = int(input("Ingrese el numero de fila a borrar:  "))
                DeleteLine = LineaDelete + 1
                hoja_Productos.delete_rows(DeleteLine, 1)
                remove("Inventario.xlsx")
                libro.save("Inventario.xlsx")
                Menu_Productos = input("Seleccione la opcion deseada:\n i.   Agregar producto\n ii.  Editar producto\n iii. Eliminar producto\n iv.  Listar productos\n v.   Enviar cotizacion por correo\n vi.  Salir\n")
            if Menu_Productos == "iv":
                #listar productos
                productos = []
                diccionarioProducto = {}

                for row in range(2, hoja_Productos.max_row + 1):
                    producto = hoja_Productos["A" + str(row)].value
                    precio = hoja_Productos["B" + str(row)].value
                    existencia = hoja_Productos["C"+str(row)].value
                    diccionarioProducto['producto'] = producto
                    diccionarioProducto['precio'] = precio
                    diccionarioProducto['existencia'] = existencia
                    productos.append(diccionarioProducto)
                    diccionarioProducto = {}

                for products in productos:
                    print("Producto: ", products['producto'] )
                    print("   Precio: "+ str(products['precio'])+ " Existencia: " + str(products['existencia']))
                Menu_Productos = input("Seleccione la opcion deseada:\n i.   Agregar producto\n ii.  Editar producto\n iii. Eliminar producto\n iv.  Listar productos\n v.   Enviar cotizacion por correo\n vi.  Salir\n")
            if Menu_Productos == "v":
                clientesCotizar = []
                diccionarioClientesCotizar = {}

                for row in range(2, hoja_Clientes.max_row + 1):
                    nombreCliente = hoja_Clientes["A" + str(row)].value
                    diccionarioClientesCotizar['nombre'] = nombreCliente
                    clientesCotizar.append(diccionarioClientesCotizar)
                    diccionarioClientesCotizar = {}
                print("Los clientes son: ")
                for clientes in clientesCotizar:
                    print("El nombre del cliente es: "+clientes['nombre'])

                # Agregar codigo para cotizacion y enviar correo
                Nombre_Cliente = input("Seleccione el nombre del cliente o ingrese un nuevo nombre: \n")
                Correo_Cliente = input("Ingrese su correo electronico: \n")
                ProductosCotizar = []
                diccionarioProductoCotizar = {}

                for row in range(2, hoja_Productos.max_row + 1):
                    producto = hoja_Productos["A" + str(row)].value
                    precio = hoja_Productos["B" + str(row)].value
                    diccionarioProductoCotizar['producto'] = producto
                    diccionarioProductoCotizar['precio'] = precio
                    ProductosCotizar.append(diccionarioProductoCotizar)
                    diccionarioProductoCotizar = {}

                Iterador = 1
                for cotizar in ProductosCotizar:
                    print("Numero de linea: "+str(Iterador)+" Producto: "+cotizar['producto']+" Precio: "+str(cotizar['precio']))
                    Iterador += 1

                Producto_Cotizar = int(input("Ingrese el numero de la linea del producto que desea cotizar: \n"))
                CotizarAumentado = Producto_Cotizar + 1  
                Precio_Cotizado = str(hoja_Productos['B' + str(CotizarAumentado)].value)
                #Creacion del documento "cotizacion.docx"
                Doc = Document()
                Doc.add_heading("COTIZACION DE PRODUCTOS")
                Doc.add_paragraph("Estimado "+Nombre_Cliente+"\n El precio de nuestro producto "+hoja_Productos['A' +str(CotizarAumentado)].value+"\n Es "+Precio_Cotizado)
                Doc.save("Cotizacion.docx")  
                mensaje = EmailMessage()
                mensaje["Subject"] = ("Hola "+Nombre_Cliente+" Como Estas")
                mensaje["From"]= DIRECCION_DE_ORIGEN
                mensaje["To"]= Correo_Cliente

                mensaje.set_content("Este es el cuerpo del mensaje")

                mensaje.add_alternative(""" 
                    <p>
                        <h1>Formato de Cotizacion</h1>
                        <h3>Te hemos enviado un documento de word con el producto que cotizaste</h3>
                        <h3>Buen Dia</h3>
                    </p>              
                """, subtype = "html")

                NombreDelArchivo = "Cotizacion.docx"
                ctype, encoding = mimetypes.guess_type(NombreDelArchivo)
                if ctype is None or encoding is not None:
                    ctype = 'application/octet-stream'

                tipoPrincipal, subTipo = ctype.split('/', 1)
                maintype, subtype = ctype.split('/', 1)
                
                with open(NombreDelArchivo, 'rb') as archivoLeido:
                    mensaje.add_attachment(archivoLeido.read(), maintype = tipoPrincipal, subtype = subTipo, filename = NombreDelArchivo)
                context = ssl.create_default_context()

                smtp = smtplib.SMTP(DIRECCION_DEL_SERVIDOR, PUERTO)
                smtp.starttls()
                smtp.login(DIRECCION_DE_ORIGEN, CONTRASENA)
                smtp.send_message(mensaje)
            
                Menu_Productos = input("Seleccione la opcion deseada:\n i.   Agregar producto\n ii.  Editar producto\n iii. Eliminar producto\n iv.  Listar productos\n v.   Enviar cotizacion por correo\n vi.  Salir\n")
        Menu_Principal = input("Seleccione la opcion deseada:\n a. Productos\n b. Clientes\n c. Pedidos\n d. Informes\n e. Varios\n f. Salir\n")
    if Menu_Principal == "b":
        Menu_Clientes = input("Seleccione la opcion deseada:\n i.   Agregar cliente\n ii.  Editar cliente\n iii. Eliminar cliente\n iv.  Listar cliente\n v.   Salir\n")
        while Menu_Clientes != "v":
            if Menu_Clientes == "i":
                # Agregar clientes
                hoja_Clientes['A1'].value = "Nombre"
                hoja_Clientes['B1'].value = "NIT"
                hoja_Clientes['C1'].value = "Direccion"
                agregarClientes = []
                AgregarClientesDicctionary = {}
                NombreClienteAgg = input("Ingrese el nombre del cliente:\n")
                AgregarClientesDicctionary['nombrecliente'] = NombreClienteAgg
                NITCLienteAGG = int(input("Ingrese el nit del cliente:\n"))
                AgregarClientesDicctionary['NITcliente'] = NITCLienteAGG
                DireccionClienteAGG = input("Ingrese la direccion del cliente:\n")
                AgregarClientesDicctionary['DireccionCLiente'] = DireccionClienteAGG
                agregarClientes.append(AgregarClientesDicctionary)
                
                primeraFilaCliente = hoja_Clientes.max_row + 1
                for Client in agregarClientes:
                    hoja_Clientes['A' + str(primeraFilaCliente)].value = Client['nombrecliente']
                    hoja_Clientes['B' + str(primeraFilaCliente)].value = Client['NITcliente']
                    hoja_Clientes['C' + str(primeraFilaCliente)].value = Client['DireccionCLiente']
                remove("Inventario.xlsx")
                libro.save("Inventario.xlsx")
                Menu_Clientes = input("Seleccione la opcion deseada:\n i.   Agregar cliente\n ii.  Editar cliente\n iii. Eliminar cliente\n iv.  Listar cliente\n v.   Salir\n")
            if Menu_Clientes == "ii":
                #codigo editar clientes
                ClienteEdit = []
                diccionarioClienteEdit = {}

                for client in range(2, hoja_Clientes.max_row + 1):
                    nombreClienteEdit = hoja_Clientes["A" + str(client)].value
                    NitClienteEdit = hoja_Clientes["B" + str(client)].value
                    DireccionEdit = hoja_Clientes["C"+str(client)].value
                    diccionarioClienteEdit['nombre'] = nombreClienteEdit
                    diccionarioClienteEdit['nit'] = NitClienteEdit
                    diccionarioClienteEdit['direccion'] = DireccionEdit
                    ClienteEdit.append(diccionarioClienteEdit)
                    diccionarioClienteEdit = {}
                IteradorClienteEdit = 1
                for ClientEdit in ClienteEdit:
                    print("Linea: " + str(IteradorClienteEdit) +" Nombre: " + ClientEdit['nombre'] +" NIT: " + str(ClientEdit['nit']) + " Direccion: " + str(ClientEdit['direccion']))
                    IteradorClienteEdit += 1
                EditarSelect = int(input("Ingrese el numero de linea a editar:  "))
                Edit = EditarSelect + 1
                hoja_Clientes['A1'].value = "Nombre"
                hoja_Clientes['B1'].value = "NIT"
                hoja_Clientes['C1'].value = "Direccion"
                agregarClientEdit = []
                AgregarDicctionaryClientEdit = {}
                ClientAggEdit = input("Ingrese el nuevo nombre del cliente: "+ hoja_Clientes['A' + str(Edit)].value + "\n")
                AgregarDicctionaryClientEdit['nombre'] = ClientAggEdit
                NitAggEdit = int(input("Ingrese el nuevo Nit del cliente:  "+ str(hoja_Clientes['B' + str(Edit)].value) + "\n"))
                AgregarDicctionaryClientEdit['nit'] = NitAggEdit
                DirecciAggEdit = input("Ingrese la  nueva direccion del cliente:  "+ hoja_Clientes['C' + str(Edit)].value +"\n")
                AgregarDicctionaryClientEdit['direccion'] = DirecciAggEdit
                agregarClientEdit.append(AgregarDicctionaryClientEdit)
                
                for ClientEdit in agregarClientEdit:
                    hoja_Clientes['A' + str(Edit)].value = ClientEdit['nombre']
                    hoja_Clientes['B' + str(Edit)].value = ClientEdit['nit']
                    hoja_Clientes['C' + str(Edit)].value = ClientEdit['direccion']
                remove("Inventario.xlsx")
                libro.save("Inventario.xlsx")
                Menu_Clientes = input("Seleccione la opcion deseada:\n i.   Agregar cliente\n ii.  Editar cliente\n iii. Eliminar cliente\n iv.  Listar cliente\n v.   Salir\n")
            if Menu_Clientes == "iii":
                #Agregar codigo eliminar clientes
                ClienteDelete = []
                diccionarioClienteDelete = {}

                for client in range(2, hoja_Clientes.max_row + 1):
                    nombreClienteEdit = hoja_Clientes["A" + str(client)].value
                    NitClienteEdit = hoja_Clientes["B" + str(client)].value
                    DireccionEdit = hoja_Clientes["C"+str(client)].value
                    diccionarioClienteDelete['nombre'] = nombreClienteEdit
                    diccionarioClienteDelete['nit'] = NitClienteEdit
                    diccionarioClienteDelete['direccion'] = DireccionEdit
                    ClienteDelete.append(diccionarioClienteDelete)
                    diccionarioClienteDelete = {}
                IteradorClienteEdit = 1
                for ClientEdit in ClienteDelete:
                    print("Linea: " + str(IteradorClienteEdit) +" Nombre: " + ClientEdit['nombre'] +" NIT: " + str(ClientEdit['nit']) + " Direccion: " + str(ClientEdit['direccion']))
                    IteradorClienteEdit += 1
                LineaDelete = int(input("Ingrese el numero de fila a borrar:  "))
                DeleteLine = LineaDelete + 1
                hoja_Clientes.delete_rows(DeleteLine, 1)
                remove("Inventario.xlsx")
                libro.save("Inventario.xlsx")
                Menu_Clientes = input("Seleccione la opcion deseada:\n i.   Agregar cliente\n ii.  Editar cliente\n iii. Eliminar cliente\n iv.  Listar cliente\n v.   Salir\n")
            if Menu_Clientes == "iv":
                #Listar clientes
                clientes = []
                diccionarioClientes = {}

                for row in range(2, hoja_Clientes.max_row + 1):
                    nombreCliente = hoja_Clientes["A" + str(row)].value
                    telefonoCliente = hoja_Clientes["B" + str(row)].value
                    DireccionCliente = hoja_Clientes["C"+str(row)].value
                    diccionarioClientes['nombre'] = nombreCliente
                    diccionarioClientes['nit'] = telefonoCliente
                    diccionarioClientes['direccion'] = DireccionCliente
                    clientes.append(diccionarioClientes)
                    diccionarioClientes = {}
                for client in clientes:
                    print("Nombre del cliente: "+client['nombre'])
                    print("   NIT: "+ str(client['nit'])+ " Dirección: " + str(client['direccion']))
                Menu_Clientes = input("Seleccione la opcion deseada:\n i.   Agregar cliente\n ii.  Editar cliente\n iii. Eliminar cliente\n iv.  Listar cliente\n v.   Salir\n")
        Menu_Principal = input("Seleccione la opcion deseada:\n a. Productos\n b. Clientes\n c. Pedidos\n d. Informes\n e. Varios\n f. Salir\n")
    if Menu_Principal == "c":
        menu_Pedidos = input("Seleccione la opcion desesada:\n i. Agregar pedido\n ii. Eliminar pedido\n iii. Listar pedido\n iv. Salir\n")
        while menu_Pedidos != "iv":
            if menu_Pedidos == "i":
                clientesPedidos = []
                diccionarioClientesPedidos = {}

                for row in range(2, hoja_Clientes.max_row + 1):
                    nombreCliente = hoja_Clientes["A" + str(row)].value
                    diccionarioClientesPedidos['nombre'] = nombreCliente
                    clientesPedidos.append(diccionarioClientesPedidos)
                    diccionarioClientesPedidos = {}
                print("Los clientes son: ")
                for clientes in clientesPedidos:
                    print("El nombre del cliente es: "+clientes['nombre'])
                #Agregar codigo agregar pedido
                hoja_Pedidos['A1'].value = "Cliente"
                hoja_Pedidos['B1'].value = "Producto"
                hoja_Pedidos['C1'].value = "Cantidad"
                hoja_Pedidos['D1'].value = "Valor pedido"
                agregarPedido = []
                AgregarPedidoDicctionary = {}
                NombreClienteAgg = input("Ingrese uno de los nombres que se muestran en la lista:\n")
                AgregarPedidoDicctionary['nombrecliente'] = NombreClienteAgg
                ProductosPedido = []
                diccionarioProductoPedido = {}

                for row in range(2, hoja_Productos.max_row + 1):
                    producto = hoja_Productos["A" + str(row)].value
                    diccionarioProductoPedido['producto'] = producto
                    ProductosPedido.append(diccionarioProductoPedido)
                    diccionarioProductoPedido = {}

                Iterador = 1
                for cotizar in ProductosPedido:
                    print("Numero de linea: "+str(Iterador)+" Producto: "+cotizar['producto'])
                    Iterador += 1
                NombreproductoPedido = int(input("Ingrese el numero de la linea del producto\n"))
                Lineapedido = NombreproductoPedido + 1
                AgregarPedidoDicctionary['nombreproducto'] = hoja_Productos['A' + str(Lineapedido)].value
                CantidadDesada = int(input("Ingrese la cantidad deseada: \n"))
                AgregarPedidoDicctionary['cantidad'] = CantidadDesada 
                AgregarPedidoDicctionary['valor'] = CantidadDesada * hoja_Productos['B'+str(Lineapedido)].value
                agregarPedido.append(AgregarPedidoDicctionary)
                
                
                primeraFilpedido = hoja_Pedidos.max_row + 1
                for Pedido in agregarPedido:
                    hoja_Pedidos['A' + str(primeraFilpedido)].value = Pedido['nombrecliente']
                    hoja_Pedidos['B' + str(primeraFilpedido)].value = Pedido['nombreproducto']
                    hoja_Pedidos['C' + str(primeraFilpedido)].value = Pedido['cantidad']
                    hoja_Pedidos['D' + str(primeraFilpedido)].value = Pedido['valor']
                ModificacionCantidad = hoja_Productos['C'+str(Lineapedido)].value - CantidadDesada
                hoja_Productos['C'+str(Lineapedido)]= ModificacionCantidad
                remove("Inventario.xlsx")
                libro.save("Inventario.xlsx")
                menu_Pedidos = input("Seleccione la opcion desesada:\n i. Agregar pedido\n ii. Eliminar pedido\n iii. Listar pedido\n iv. Salir\n")
            if menu_Pedidos == "ii":
                #Agregar codigo eliminar pedido
                pedidosdelte = []
                diccionarioPedidosdelt = {}

                for row in range(2, hoja_Pedidos.max_row + 1):
                    NombreClientePed = hoja_Pedidos["A" + str(row)].value
                    NombreProductosPed = hoja_Pedidos["B" + str(row)].value
                    CantidadPedido = hoja_Pedidos["C" + str(row)].value
                    valorPedido = hoja_Pedidos["D"+str(row)].value
                    diccionarioPedidosdelt['nombreCliente'] = NombreClientePed
                    diccionarioPedidosdelt['nombreProdcuto'] = NombreProductosPed
                    diccionarioPedidosdelt['cantidadPedido'] = CantidadPedido
                    diccionarioPedidosdelt['valorPedido'] = valorPedido
                    pedidosdelte.append(diccionarioPedidosdelt)
                    diccionarioPedidosdelt = {}
                iteradordelteped = 1
                for Deltclient in pedidosdelte:
                    print("Linea: " + str(iteradordelteped)+"  Cliente: "+Deltclient['nombreCliente']+"  Producto: "+ Deltclient['nombreProdcuto']+ "  Cantidad: " + str(Deltclient['cantidadPedido'])+"  Valor pedido: "+str(Deltclient['valorPedido']))
                    iteradordelteped += 1
                LineaDelete = int(input("Ingrese el numero de fila a borrar:  "))
                DeleteLine = LineaDelete + 1
                hoja_Pedidos.delete_rows(DeleteLine, 1)
                remove("Inventario.xlsx")
                libro.save("Inventario.xlsx")
                menu_Pedidos = input("Seleccione la opcion desesada:\n i. Agregar pedido\n ii. Eliminar pedido\n iii. Listar pedido\n iv. Salir\n")
            if menu_Pedidos == "iii":
                #listar pedidos
                pedidos = []
                diccionarioPedidos = {}

                for row in range(2, hoja_Pedidos.max_row + 1):
                    NombreClientePed = hoja_Pedidos["A" + str(row)].value
                    NombreProductosPed = hoja_Pedidos["B" + str(row)].value
                    CantidadPedido = hoja_Pedidos["C" + str(row)].value
                    valorPedido = hoja_Pedidos["D"+str(row)].value
                    diccionarioPedidos['nombreCliente'] = NombreClientePed
                    diccionarioPedidos['nombreProdcuto'] = NombreProductosPed
                    diccionarioPedidos['cantidadPedido'] = CantidadPedido
                    diccionarioPedidos['valorPedido'] = valorPedido
                    pedidos.append(diccionarioPedidos)
                    diccionarioPedidos = {}
                for pedido in pedidos:
                    print("Cliente: "+pedido['nombreCliente'])
                    print("  Producto: "+ pedido['nombreProdcuto']+ "  Cantidad: " + str(pedido['cantidadPedido'])+"  Valor pedido: "+str(pedido['valorPedido']))
                menu_Pedidos = input("Seleccione la opcion desesada:\n i. Agregar pedido\n ii. Eliminar pedido\n iii. Listar pedido\n iv. Salir\n")
        Menu_Principal = input("Seleccione la opcion deseada:\n a. Productos\n b. Clientes\n c. Pedidos\n d. Informes\n e. Varios\n f. Salir\n")
    if Menu_Principal == "d":
        menu_Informes = input("Selecciones la opcion deseada:\n i. Total venta por cliente\n ii. Total ventas por producto\n iii. Salir\n")
        while menu_Informes != "iii":
            if menu_Informes == "i":
                #Codigo total ventas
                pedidosClienteInf = []
                diccionarioInforme = {}
                iterador = 2
                total_ventas_clientes = 0.0
                for row in range(2, hoja_Pedidos.max_row + 1):
                    NombreClientePed = hoja_Pedidos["A" + str(row)].value
                    valorPedido = hoja_Pedidos["D"+str(row)].value
                    diccionarioInforme['nombreCliente'] = NombreClientePed
                    diccionarioInforme['valorPedido'] = valorPedido
                    pedidosClienteInf.append(diccionarioInforme)
                    diccionarioInforme = {}
                
                for infpedi in pedidosClienteInf:
                    print("Nombre cliente: "+infpedi['nombreCliente']+" total venta: "+str(infpedi['valorPedido']))
                    total_ventas_clientes += hoja_Pedidos["D"+str(iterador)].value 
                    iterador += 1
                print(total_ventas_clientes)
                menu_Informes = input("Selecciones la opcion deseada:\n i. Total venta por cliente\n ii. Total ventas por producto\n iii. Salir\n")
            if menu_Informes == "ii":
                #Codigo total ventas productos
                pedidosProdcuInf = []
                diccionarioInformeProducto = {}
                iterador = 2
                total_ventas_productos = 0.0
                for row in range(2, hoja_Pedidos.max_row + 1):
                    NombreProductPed = hoja_Pedidos["B" + str(row)].value
                    valorPedido = hoja_Pedidos["D"+str(row)].value
                    diccionarioInformeProducto['Nombreproducto'] = NombreProductPed
                    diccionarioInformeProducto['valorPedido'] = valorPedido
                    pedidosProdcuInf.append(diccionarioInformeProducto)
                    diccionarioInformeProducto = {}
                
                for infpedi in pedidosProdcuInf:
                    print("Nombre prdocuto: "+infpedi['Nombreproducto']+" total venta: "+str(infpedi['valorPedido']))
                    total_ventas_productos += hoja_Pedidos["D"+str(iterador)].value 
                    iterador += 1
                print(total_ventas_productos)
                menu_Informes = input("Selecciones la opcion deseada:\n i. Total venta por cliente\n ii. Total ventas por producto\n iii. Salir\n")
        Menu_Principal = input("Seleccione la opcion deseada:\n a. Productos\n b. Clientes\n c. Pedidos\n d. Informes\n e. Varios\n f. Salir\n")
    if Menu_Principal == "e":
       
        decisionArchivo = input(" 1.Creara un nuevo archivo de seguirdad\n 2.Borar existente y crear uno nuevo\n")
        if decisionArchivo == "1":
            libro.save("Inventario_Copia_de_Seguridad.xlsx")
        if decisionArchivo == "2":
            remove("Inventario_Copia_de_Seguridad.xlsx")
            libro.save("Inventario_Copia_de_Seguridad.xlsx")
        mensaje = EmailMessage()
        mensaje["Subject"] = "Hola buen dia señor coordinador"
        mensaje["From"] = DIRECCION_DE_ORIGEN
        mensaje["To"] = input("Ingrese su correo\n")

        mensaje.set_content("Este es el cuerpo del mensaje")

        mensaje.add_alternative("""
                    <p>
                        <h1>ADJUNTAMOS COPIA DE SEGURIDAD</h1>
                        <h3>Te enviamos esta copia de seguridad para respaldar los datos del Inventario</h3>
                    </p>               
                """, subtype = "html")
        nombre_de_archivo = "Inventario_Copia_de_Seguridad.xlsx"
        ctype, encoding = mimetypes.guess_type(nombre_de_archivo)
        if ctype is None or encoding is not None:
            ctype = 'application/octet-stream'

        tipoPrincipal, subTipo = ctype.split('/', 1)

        with open(nombre_de_archivo, 'rb') as archivoLeido:
            mensaje.add_attachment(archivoLeido.read(), maintype=tipoPrincipal, subtype = subTipo, filename = nombre_de_archivo)

        context = ssl.create_default_context()

        smtp = smtplib.SMTP(DIRECCION_DEL_SERVIDOR, PUERTO)
        smtp.starttls()
        smtp.login(DIRECCION_DE_ORIGEN, CONTRASENA)
        smtp.send_message(mensaje)

        Menu_Principal = input("Seleccione la opcion deseada:\n a. Productos\n b. Clientes\n c. Pedidos\n d. Informes\n e. Varios\n f. Salir\n")
